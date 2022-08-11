# 书写人：胡高远
# 书写时间：2022/7/19  16:15
from docx import Document
from docx.shared import Cm
doc = Document()

# 添加图片
# doc.add_picture('logo.png')
doc.add_picture('logo.png',width = Cm(5))
doc.save('我的文档2.docx')