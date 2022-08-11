# 书写人：胡高远
# 书写时间：2022/7/19  16:55
from docx import Document
import datetime
doc = Document()

'''添加标题'''
doc.add_heading('请假条',level=1)
para1 = doc.add_paragraph('尊进的领导：')

# 以下内容从工作台输入
reason = input('请输入请假原因：')
start_date = input('请输入开始日期：')
end_date = input('请输入结束日期：')

# 将字符串类型的日期转成日期类型
start_date = datetime.datetime.strptime(start_date,'%Y-%m-%d').date()
end_date = datetime.datetime.strptime(end_date,'%Y-%m-%d').date()

# 计算两个日期之差
s_date = str(end_date-start_date).split(' ')[0]

# 构建格式化字符串
s = '因{0}需要请假，请假时间从{1}到{2}，一共{3}天，恳请领导批准。'.format(reason,start_date,end_date,s_date)

# 将构建的字符串添加到文档的段落中
doc.add_paragraph(s)
person = input('请输入请假人：')
doc.add_paragraph('请假人：'+person)

# 请假的时间为系统的当前时间
date = datetime.datetime.now().strftime('%Y-%m-%d')

doc.add_paragraph(date)
doc.save('请假条.docx')
