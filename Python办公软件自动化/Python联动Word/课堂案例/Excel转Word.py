# 书写人：胡高远
# 书写时间：2022/7/21  14:35

import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


workbook = openpyxl.load_workbook('成绩表.xlsx')
sheet = workbook['Sheet1']
rows = sheet.rows
lst = []
for row in rows:
    row_lst = []
    for cell in row:
        row_lst.append(cell.value)
    lst.append(row_lst)
# print(lst)

# 新建document对象
doc = Document()
title = doc.add_heading(lst[0][0],level=1)
for run in title.runs:
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.size = Pt(14)
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'),'微软雅黑')

title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 画表格
table = doc.add_table(rows=len(lst)-1,cols=len(lst[1]),style='Light Shading Accent 3')
for index in range(0,len(lst)-1):
    cells = table.rows[index].cells
    for idx in range(0,len(lst[1])):
        if str(lst[index+1][idx])!='None':
            cells[idx].text = str(lst[index+1][idx])
            cells[idx].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

doc.save('成绩表.docx')