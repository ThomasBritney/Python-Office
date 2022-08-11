# 书写人：胡高远
# 书写时间：2022/7/20  15:28

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 读取内容
with open('学如弓弩.txt','r',encoding = 'utf-8') as file:
    line = file.readlines()
    for item in line:
        if item == '\n':
            line.remove(item)
# print(len(line))

#向Word文档写入内容
doc = Document()
title = doc.add_heading(line[0],level=1)
for run in title.runs:
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.size = Pt(20)
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'),'微软雅黑')
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 设置日期的右对齐
para1 = doc.add_paragraph(line[1])
para1.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT

# 设置正文部分
for item in range(2,len(line)):
    para = doc.add_paragraph(line[item])
    for run in para.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(12)
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'),'微软雅黑')
    para.paragraph_format.line_spacing = 2.0


doc.save('学如弓弩.docx')
