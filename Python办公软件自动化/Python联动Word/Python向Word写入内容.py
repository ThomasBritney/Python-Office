# 书写人：胡高远
# 书写时间：2022/7/19  15:58
from docx import Document
doc = Document()    #创建Document对象

# 添加标题
doc.add_heading('晓出净慈寺送林子方',level=1)

# 添加段落
para1 = doc.add_paragraph('作者：杨万里')
para2 = doc.add_paragraph('毕竟西湖六月中，风光不与四时同。')

# 向段落添加文字块
para3 = doc.add_paragraph()
para3.add_run('接天莲叶无穷碧，').font.bold = True
para3.add_run('映日荷花别样红。').font.italic = True
para3.add_run('小荷才露尖尖角，早有蜻蜓立上头')

# 添加分页
doc.add_page_break()

# 保存成Word文件
doc.save('杨万里.docx')
