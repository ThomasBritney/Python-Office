# 书写人：胡高远
# 书写时间：2022/7/20  14:55
from docx import *
from docx.enum.style import WD_STYLE_TYPE
doc = Document()
styles = doc.styles
for s in styles:
    if s.type ==WD_STYLE_TYPE.TABLE:
        doc.add_paragraph('表格样式：'+s.name)
        #向Word文档中添加表格
        table = doc.add_table(3,3,style=s)
        cell = table.rows[0].cells
        cell[0].text = '第一列'
        cell[1].text = '第二列'
        cell[2].text = '第三列'
        doc.add_paragraph('\n')

doc.save('表格样式.docx')


